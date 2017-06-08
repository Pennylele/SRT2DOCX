#! /usr/bin/env python3
# SRT2DOCX.py is used for writing content from a SRT file into a DOCX file.

from docx import Document
from docx.text.run import Font, Run
import io, os, re
import docx
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

cwd = os.getcwd()


def func(outputFiles):
    for file in os.listdir(cwd):
        if file.endswith('.srt'):
            f = io.open(file, encoding="latin_1", mode="r+")
            regexContent = re.compile('\d{2}:\d{2}:\d{2}(?=,\d{3}\s-->)')
            timeCodes = regexContent.findall(f.read())
            print(timeCodes)
            document = Document()
            #document.add_heading('SRT_converted', 0)


            f = io.open(file, encoding="latin_1", mode="r+")
            regexContent2 = re.compile('(?<=\d{2}:\d{2}:\d{2},\d{3}\s-->\s\d{2}:\d{2}:\d{2},\d{3}\n).*')
            subtitles = regexContent2.findall(f.read())
            print(subtitles)
            dictionary = dict(zip(timeCodes, subtitles))
            print(dictionary)

            table = document.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            # hdr_cells[0].text = 'TimeCode'
            cell1 = hdr_cells[0]
            cell1.text = 'TimeCode'
            run = cell1.paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor(255, 255, 255)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
            table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
            # hdr_cells[1].text = 'Personnel'
            cell2 = hdr_cells[1]
            cell2.text = 'Personnel'
            run = cell2.paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor(255, 255, 255)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
            table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
            #hdr_cells[2].text = 'Original Content Transcription'
            cell2 = hdr_cells[2]
            cell2.text = 'Original Content Transcription'
            run = cell2.paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor(255, 255, 255)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
            table.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
            #hdr_cells[3].text = 'Translation'
            cell3 = hdr_cells[3]
            cell3.text = 'Translation'
            run = cell2.paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor(255, 255, 255)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
            table.rows[0].cells[3]._tc.get_or_add_tcPr().append(shading_elm_1)

            for timeCode, subtitle in sorted(dictionary.items()):
                row_cells = table.add_row().cells
                row_cells[0].text = str(timeCode)
                row_cells[2].text = str(subtitle)
                row_cells[3].text = str(subtitle)
            document.save(file+'.docx')
            for file in os.listdir(cwd):
                filename = os.path.join(cwd, file)
                #oldbase = os.path.splitext(filename)
                newname = filename.replace('.srt.docx', '.docx')
                os.rename(filename, newname)



outputFiles = os.listdir(cwd)
func(outputFiles)
